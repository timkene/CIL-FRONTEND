"""
Vercel Python Serverless Function
POST /api/sla_generate
Generates a filled SLA .docx, uploads to Supabase Storage, logs to sla_documents.
"""

from http.server import BaseHTTPRequestHandler
import json, copy, io, re, os
from datetime import datetime, timedelta, timezone
from docx import Document
from supabase import create_client

# ── Config ────────────────────────────────────────────────────────────────────
HERE          = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(HERE, "clearline_sla_template.docx")

SUPABASE_URL  = os.environ.get("SUPABASE_URL", "https://zxxkcvkrpdsvfljrjgqy.supabase.co")
SUPABASE_KEY  = os.environ.get("SUPABASE_KEY", "")

MONTHS  = ["January","February","March","April","May","June",
           "July","August","September","October","November","December"]
ORDINAL = {
    1:"1st",2:"2nd",3:"3rd",4:"4th",5:"5th",6:"6th",7:"7th",
    8:"8th",9:"9th",10:"10th",11:"11th",12:"12th",13:"13th",
    14:"14th",15:"15th",16:"16th",17:"17th",18:"18th",19:"19th",
    20:"20th",21:"21st",22:"22nd",23:"23rd",24:"24th",25:"25th",
    26:"26th",27:"27th",28:"28th",29:"29th",30:"30th",31:"31st",
}
DOTS = r"[…\.]{2,}"


# ── Document generation ───────────────────────────────────────────────────────
def _fmt_naira(value) -> str:
    try:
        n = float(str(value).replace(",","").replace("₦","").replace("N","").strip())
        return f"₦{int(n):,}" if n == int(n) else f"₦{n:,.2f}"
    except Exception:
        return str(value)

def _set_cell(row, col_idx, text):
    cell = row.cells[col_idx]
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = text
        else:
            para.add_run(text)

def generate_sla_bytes(data: dict) -> bytes:
    doc   = Document(TEMPLATE_PATH)
    paras = doc.paragraphs
    name  = data["company_name"].strip().upper()

    # 1. Cover page — para 25
    p25 = paras[25]
    for r in p25.runs: r.text = ""
    if p25.runs: p25.runs[0].text = name

    # 2. Opening paragraph — para 31
    p31 = paras[31]
    for run in p31.runs:
        if re.search(DOTS, run.text) and "day of" in run.text:
            run.text = re.sub(DOTS, ORDINAL[data["contract_day"]], run.text, count=1)
            run.text = re.sub(DOTS, data["contract_month"], run.text, count=1)
            break
    for run in p31.runs:
        t = run.text.strip()
        if re.fullmatch(r"[…\.\s]+", t) or (re.search(DOTS, t) and len(t) < 60 and "address" not in t):
            run.text = " " + name + " "; break
    for run in p31.runs:
        if "address is at" in run.text:
            run.text = re.sub(DOTS, data["company_address"], run.text, count=1); break

    # 3. Premium paragraph — para 56
    p56  = paras[56]
    full = "".join(r.text for r in p56.runs)
    full = re.sub(DOTS, data["premium_naira"],     full, count=1)
    full = re.sub(DOTS, data["premium_words"],     full, count=1)
    full = re.sub(DOTS, data["num_beneficiaries"], full, count=1)
    for r in p56.runs: r.text = ""
    if p56.runs: p56.runs[0].text = full

    # 4. Plans table
    table    = doc.tables[0]
    data_row = table.rows[1]
    total_tr = table.rows[2]._tr
    plans        = data["plans"]
    total_lives  = sum(int(str(p.get("num_lives",0)).replace(",","") or 0) for p in plans)
    total_amount = sum(float(str(p.get("amount",0)).replace(",","").replace("₦","") or 0) for p in plans)
    first = plans[0] if plans else {}
    _set_cell(data_row, 0, "1.")
    _set_cell(data_row, 1, first.get("plan_type",""))
    _set_cell(data_row, 2, first.get("description",""))
    _set_cell(data_row, 3, str(first.get("num_lives","")))
    a = first.get("amount","")
    _set_cell(data_row, 4, _fmt_naira(a) if a else "")
    for idx, plan in enumerate(plans[1:], start=2):
        new_tr = copy.deepcopy(data_row._tr)
        total_tr.addprevious(new_tr)
        new_row = table.rows[idx]
        _set_cell(new_row, 0, f"{idx}.")
        _set_cell(new_row, 1, plan.get("plan_type",""))
        _set_cell(new_row, 2, plan.get("description",""))
        _set_cell(new_row, 3, str(plan.get("num_lives","")))
        aa = plan.get("amount","")
        _set_cell(new_row, 4, _fmt_naira(aa) if aa else "")
    total_row = table.rows[-1]
    _set_cell(total_row, 2, f"TOTAL  ({total_lives:,} lives)")
    grand = int(total_amount) if total_amount == int(total_amount) else total_amount
    _set_cell(total_row, 4, _fmt_naira(grand))

    # 5. Period of cover — para 123
    p123 = paras[123]
    full = "".join(r.text for r in p123.runs)
    full = re.sub(DOTS, ORDINAL[data["start_day"]],                                full, count=1)
    full = re.sub(DOTS, data["start_month"],                                       full, count=1)
    full = re.sub(DOTS, f"{ORDINAL[data['end_day']]} {data['end_month']}",         full, count=1)
    full = full.replace("2026", data["start_year"], 1)
    full = full.replace("2027", data["end_year"],   1)
    for r in p123.runs: r.text = ""
    if p123.runs: p123.runs[0].text = full

    # 6. Signature page — para 268
    p268 = paras[268]
    for r in p268.runs: r.text = ""
    if p268.runs: p268.runs[0].text = name

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Supabase helpers ──────────────────────────────────────────────────────────
def upload_and_record(docx_bytes: bytes, data: dict) -> dict:
    sb   = create_client(SUPABASE_URL, SUPABASE_KEY)
    name = data["company_name"].strip().upper()
    slug = re.sub(r"[^\w]", "_", name)[:40]
    ts   = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    path = f"{data['start_year']}/{slug}_{ts}.docx"

    sb.storage.from_("sla-documents").upload(
        path, docx_bytes,
        {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    )
    signed = sb.storage.from_("sla-documents").create_signed_url(path, 3650 * 24 * 3600)
    url    = signed.get("signedURL") or signed.get("signed_url") or signed["data"]["signedURL"]

    month_num = lambda m: MONTHS.index(m) + 1
    now = datetime.now(timezone.utc)
    row = {
        "company_name":  name,
        "contract_start": f"{data['start_year']}-{month_num(data['start_month']):02d}-{data['start_day']:02d}",
        "contract_end":   f"{data['end_year']}-{month_num(data['end_month']):02d}-{data['end_day']:02d}",
        "generated_by":  data.get("generated_by", ""),
        "action":        "download",
        "storage_path":  path,
        "download_url":  url,
        "created_at":    now.isoformat(),
        "expires_at":    (now + timedelta(days=365)).isoformat(),
    }
    result = sb.table("sla_documents").insert(row).execute()
    record = result.data[0] if result.data else {}
    return {"id": record.get("id"), "company_name": name, "download_url": url, "storage_path": path}


# ── Vercel handler ────────────────────────────────────────────────────────────
class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        length = int(self.headers.get("Content-Length", 0))
        body   = json.loads(self.rfile.read(length))
        try:
            docx_bytes = generate_sla_bytes(body)
            result     = upload_and_record(docx_bytes, body)
            status     = 200
        except Exception as exc:
            result = {"error": str(exc)}
            status = 500
        payload = json.dumps(result).encode()
        self.send_response(status)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(payload)))
        self.end_headers()
        self.wfile.write(payload)

    def log_message(self, *args):
        pass
